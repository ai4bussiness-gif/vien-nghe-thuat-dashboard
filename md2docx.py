#!/usr/bin/env python
"""Convert ART HAU STUDIO report from Markdown to professional DOCX."""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Brand Colours ───────────────────────────────────────────────
DARK_BG     = RGBColor(0x1A, 0x1A, 0x2E)  # deep navy
ACCENT_GOLD = RGBColor(0xC9, 0xA9, 0x6E)  # gold
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT   = RGBColor(0x1A, 0x1A, 0x2E)
BODY_TEXT   = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY  = RGBColor(0xF5, 0xF3, 0xEE)
TABLE_HEAD  = RGBColor(0x1A, 0x1A, 0x2E)
TABLE_ALT   = RGBColor(0xF0, 0xEE, 0xE8)
GRAY_TEXT   = RGBColor(0x88, 0x88, 0x88)

# ─── Helpers ──────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "C9A96E")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False,
                            color=None, size=None, alignment=None, space_before=None,
                            space_after=None, font_name='Times New Roman'):
    """Add a paragraph with formatting options."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if alignment:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point with optional bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        r.font.color.rgb = BODY_TEXT
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.color.rgb = BODY_TEXT
    return p

def add_section_heading(doc, text, level=1):
    """Add a styled section heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
    return heading

def parse_inline_bold(text):
    """Parse **bold** markers into runs. Returns list of (text, is_bold, is_italic)."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    runs_data = []
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            runs_data.append((part[2:-2], True, False))
        else:
            # Also handle *italic*
            italic_parts = re.split(r'(\*.*?\*)', part)
            for ip in italic_parts:
                if ip.startswith('*') and ip.endswith('*') and not ip.startswith('**'):
                    runs_data.append((ip[1:-1], False, True))
                else:
                    runs_data.append((ip, False, False))
    return runs_data

def add_rich_paragraph(doc, text, bold=False, italic=False, color=None,
                       size=None, alignment=None, space_before=None,
                       space_after=None, font_name='Times New Roman',
                       first_line_indent=None):
    """Add paragraph with **bold** and *italic* inline parsing."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)

    # Parse inline **bold** and *italic*
    parts = parse_inline_bold(text)
    for part_text, is_bold, is_italic in parts:
        run = p.add_run(part_text)
        run.font.name = font_name
        run.font.size = Pt(size or 11)
        run.font.color.rgb = color or BODY_TEXT
        run.bold = bold or is_bold
        run.italic = italic or is_italic
    return p


# ─── Parse MD Tables ──────────────────────────────────────────────

def parse_md_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (table_data, end_idx)."""
    header_line = lines[start_idx].strip()
    if not header_line.startswith('|'):
        return None, start_idx
    headers = [h.strip() for h in header_line.split('|')[1:-1]]
    
    # Skip separator line (|---|)
    idx = start_idx + 1
    while idx < len(lines) and lines[idx].strip().startswith('|') and all(c in '|-:' for c in lines[idx].strip()):
        idx += 1
    
    rows = []
    while idx < len(lines):
        line = lines[idx].strip()
        if not line.startswith('|'):
            break
        cells = [c.strip() for c in line.split('|')[1:-1]]
        # Pad or trim to match header count
        while len(cells) < len(headers):
            cells.append('')
        cells = cells[:len(headers)]
        rows.append(cells)
        idx += 1
    
    return {'headers': headers, 'rows': rows}, idx


# ─── Build DOCX ──────────────────────────────────────────────────

def build_docx(md_path, output_path):
    doc = Document()
    
    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # ── Default Style ──
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_TEXT
    style.paragraph_format.line_spacing = 1.3
    
    # Heading styles
    for i in [1, 2, 3]:
        hs = doc.styles[f'Heading {i}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = DARK_BG
    
    # ═══════════════════ COVER PAGE ═══════════════════════════════
    # Spacer
    for _ in range(6):
        add_formatted_paragraph(doc, '', size=11)
    
    # Decorative line
    add_formatted_paragraph(doc, '━' * 50, size=10, color=ACCENT_GOLD,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_formatted_paragraph(doc, 'BÁO CÁO KẾ HOẠCH HOẠT ĐỘNG', 
                           bold=True, size=26, color=DARK_BG,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=12)
    
    add_formatted_paragraph(doc, 'ART HAU STUDIO', 
                           bold=True, size=20, color=ACCENT_GOLD,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=6)
    
    add_formatted_paragraph(doc, 'Trung tâm Nghệ thuật Trực thuộc Viện Nghệ Thuật\nĐại Học Kiến Trúc Hà Nội', 
                           italic=True, size=13, color=GRAY_TEXT,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=4)
    
    add_formatted_paragraph(doc, '─' * 40, size=10, color=ACCENT_GOLD,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)
    
    add_formatted_paragraph(doc, 'Giai đoạn: Tháng 8/2026 – Tháng 7/2027', 
                           size=12, color=DARK_BG,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=8)
    
    # Spacer
    for _ in range(4):
        add_formatted_paragraph(doc, '', size=11)
    
    add_formatted_paragraph(doc, '━' * 50, size=10, color=ACCENT_GOLD,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # ── Page break ──
    doc.add_page_break()
    
    # ═══════════════════ CONTENT ══════════════════════════════════
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Skip cover heading (lines 1-5 of md)
    idx = 0
    in_table = False
    in_list = False
    
    while idx < len(lines):
        line = lines[idx].rstrip()
        
        # Skip the MD title, subtitle
        if idx < 8:
            idx += 1
            continue
        
        # Horizontal rule
        if line.strip() in ('---', '___', '***') and len(line.strip()) >= 3:
            add_formatted_paragraph(doc, '─' * 60, size=8, color=ACCENT_GOLD,
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_before=6, space_after=6)
            idx += 1
            continue
        
        # Blockquote
        if line.startswith('> '):
            text = line[2:]
            add_rich_paragraph(doc, text, italic=True, color=GRAY_TEXT,
                              size=10.5, space_before=4, space_after=4)
            idx += 1
            continue
        
        # Heading
        if line.startswith('## '):
            text = line[3:].strip()
            # Skip the "MỤC LỤC" heading
            if 'MỤC LỤC' in text:
                idx += 1
                continue
            # Check if next line is table (for sections)
            next_lines = [l.strip() for l in lines[idx+1:idx+5]]
            is_section_heading = any(l.startswith('|') for l in next_lines)
            if is_section_heading:
                add_formatted_paragraph(doc, text.upper(), bold=True, size=14, 
                                       color=DARK_BG, space_before=14, space_after=4)
            else:
                add_section_heading(doc, text, level=2)
            idx += 1
            continue
        
        if line.startswith('### '):
            text = line[4:].strip()
            add_section_heading(doc, text, level=3)
            idx += 1
            continue
        
        # Table detection
        if line.startswith('|') and idx + 1 < len(lines) and lines[idx+1].strip().startswith('|'):
            table_data, end_idx = parse_md_table(lines, idx)
            if table_data:
                num_cols = len(table_data['headers'])
                num_rows = len(table_data['rows']) + 1  # +1 for header
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'
                
                # Set table width
                for row in table.rows:
                    for cell in row.cells:
                        cell.width = Cm(14.0 / max(num_cols, 1))
                        # Remove cell margins for compact look
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcMar = parse_xml(
                            f'<w:tcMar {nsdecls("w")}>'
                            f'  <w:top w:w="30" w:type="dxa"/>'
                            f'  <w:left w:w="60" w:type="dxa"/>'
                            f'  <w:bottom w:w="30" w:type="dxa"/>'
                            f'  <w:right w:w="60" w:type="dxa"/>'
                            f'</w:tcMar>'
                        )
                        tcPr.append(tcMar)
                
                # Header row
                for j, h in enumerate(table_data['headers']):
                    cell = table.rows[0].cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run(h)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = WHITE
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_cell_shading(cell, '1A1A2E')
                
                # Data rows
                for i, row in enumerate(table_data['rows']):
                    for j, cell_text in enumerate(row):
                        cell = table.rows[i + 1].cells[j]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        
                        # Parse **bold** in cell
                        parts = parse_inline_bold(cell_text)
                        for pt, pb, pi in parts:
                            run = p.add_run(pt)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(9.5)
                            run.font.color.rgb = DARK_TEXT
                            run.bold = pb
                            run.italic = pi
                        
                        # Alternate row shading
                        if i % 2 == 1:
                            set_cell_shading(cell, 'F5F3EE')
                        
                        # Gold left border
                        set_cell_border(cell, 
                            start={'val': 'single', 'sz': '8', 'color': 'C9A96E'})
                
                idx = end_idx
                add_formatted_paragraph(doc, '', size=6)  # spacer
                continue
        
        # List items
        if line.strip().startswith('- '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            parts = parse_inline_bold(text)
            for pt, pb, pi in parts:
                run = p.add_run(pt)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.color.rgb = BODY_TEXT
                run.bold = pb
                run.italic = pi
            idx += 1
            continue
        
        # Numbered list items
        match = re.match(r'^\d+\.\s+(.*)', line.strip())
        if match:
            text = match.group(1)
            p = doc.add_paragraph(style='List Number')
            parts = parse_inline_bold(text)
            for pt, pb, pi in parts:
                run = p.add_run(pt)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.color.rgb = BODY_TEXT
                run.bold = pb
                run.italic = pi
            idx += 1
            continue
        
        # Empty line
        if not line.strip():
            idx += 1
            continue
        
        # Regular paragraph
        if line.strip():
            add_rich_paragraph(doc, line.strip(), size=11, space_before=2, space_after=2,
                             first_line_indent=0)
        
        idx += 1
    
    # ── Page number footer ──
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = p.add_run()
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char_begin)
        
        run2 = p.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instr)
        
        run3 = p.add_run()
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fld_char_end)
        
        for r in [run, run2, run3]:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.font.color.rgb = GRAY_TEXT
    
    # ── Save ──
    doc.save(output_path)
    print(f'✅ DOCX saved: {output_path}')
    print(f'   File size: {__import__("os").path.getsize(output_path):,} bytes')


if __name__ == '__main__':
    build_docx(
        r'C:\Users\Dell\projects\vien-nghe-thuat\ART_HAU_STUDIO_KE_HOACH.md',
        r'C:\Users\Dell\projects\vien-nghe-thuat\ART_HAU_STUDIO_KE_HOACH_v2.docx'
    )
